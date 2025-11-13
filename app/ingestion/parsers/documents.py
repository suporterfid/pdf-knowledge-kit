from __future__ import annotations

import logging
from collections.abc import Iterable
from csv import Sniffer
from csv import reader as csv_reader
from pathlib import Path

try:  # pragma: no cover - optional dependency import
    from docx import Document  # type: ignore
except ImportError:  # pragma: no cover - optional dependency import
    Document = None  # type: ignore

try:  # pragma: no cover - optional dependency import
    from openpyxl import load_workbook  # type: ignore
except ImportError:  # pragma: no cover - optional dependency import
    load_workbook = None  # type: ignore

logger = logging.getLogger(__name__)

ParsedSegment = tuple[str, dict[str, object]]


def _read_text_with_fallbacks(path: Path, encodings: Iterable[str]) -> str:
    for enc in encodings:
        try:
            return path.read_text(encoding=enc)
        except Exception as exc:
            logger.debug("Failed to read %s with encoding %s: %s", path, enc, exc)
            continue
    return path.read_bytes().decode("utf-8", errors="ignore")


def read_txt_text(path: Path) -> list[ParsedSegment]:
    """Read a plain text file returning a single segment with metadata."""

    text = _read_text_with_fallbacks(path, ("utf-8", "utf-8-sig", "latin-1"))
    return [
        (
            text,
            {
                "mime_type": "text/plain",
                "page_number": 1,
            },
        )
    ]


def read_docx_text(path: Path) -> list[ParsedSegment]:
    """Extract text from a DOCX file including paragraphs and tables."""

    if Document is None:
        raise RuntimeError("python-docx is required to read DOCX files")

    document = Document(str(path))
    parts: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [
                cell.text.strip()
                for cell in row.cells
                if cell.text and cell.text.strip()
            ]
            if cells:
                parts.append(" | ".join(cells))

    text = "\n\n".join(parts)
    return [
        (
            text,
            {
                "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "page_number": 1,
            },
        )
    ]


def read_csv_text(path: Path) -> list[ParsedSegment]:
    """Read a CSV file and return one segment per row."""

    raw = _read_text_with_fallbacks(path, ("utf-8", "utf-8-sig", "latin-1"))
    lines = raw.splitlines()
    try:
        first_line = lines[0]
    except IndexError:
        first_line = ""
    try:
        dialect = Sniffer().sniff(first_line) if first_line else None
    except Exception:
        dialect = None

    segments: list[ParsedSegment] = []
    for idx, row in enumerate(
        csv_reader(lines, dialect=dialect) if dialect else csv_reader(lines),
        start=1,
    ):
        row_text = ", ".join(cell.strip() for cell in row if cell.strip())
        if row_text:
            segments.append(
                (
                    row_text,
                    {
                        "mime_type": "text/csv",
                        "row_number": idx,
                    },
                )
            )
    return segments or [
        (
            raw,
            {
                "mime_type": "text/csv",
            },
        )
    ]


def read_xlsx_text(path: Path) -> list[ParsedSegment]:
    """Read an XLSX workbook, returning segments per sheet row."""

    if load_workbook is None:
        raise RuntimeError("openpyxl is required to read XLSX files")

    workbook = load_workbook(filename=str(path), data_only=True, read_only=True)
    segments: list[ParsedSegment] = []

    for sheet in workbook.worksheets:
        for idx, row in enumerate(sheet.iter_rows(values_only=True), start=1):
            cells = [
                str(cell).strip()
                for cell in row
                if cell is not None and str(cell).strip()
            ]
            if cells:
                segments.append(
                    (
                        " | ".join(cells),
                        {
                            "mime_type": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                            "sheet_name": sheet.title,
                            "row_number": idx,
                        },
                    )
                )

    if not segments:
        segments.append(
            (
                "",
                {
                    "mime_type": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                },
            )
        )
    workbook.close()
    return segments
